# Importa as bibliotecas necessárias do Flask
from flask import Flask, render_template, request, redirect, url_for, send_file
from docx import Document
import mysql.connector
import pandas as pd
from reportlab.pdfgen import canvas

# Inicia a aplicação Flask
app = Flask(__name__)

# Conexão com o banco de dados MySQL
db = mysql.connector.connect(
    host="localhost",  # Host do MySQL
    user="root",  # Usuário do MySQL
    password="1312",  # Senha do MySQL
    database="total_moedas"  # Nome do banco de dados
)
@app.route('/apagar_dados', methods=['POST'])
def apagar_dados():
    try:
        # Cria um cursor para executar comandos SQL
        cursor = db.cursor()

        # Comando SQL para excluir todas as informações das tabelas
        sql_delete_pessoas = "DELETE FROM pessoas"
        sql_delete_moedas = "DELETE FROM moedas"

        # Executa os comandos SQL
        cursor.execute(sql_delete_pessoas)
        cursor.execute(sql_delete_moedas)

        # Confirma as alterações no banco de dados
        db.commit()

        # Fecha o cursor
        cursor.close()

        # Redireciona para a página inicial
        return redirect(url_for('index'))

    except Exception as e:
        # Em caso de erro, desfaz as alterações e exibe uma mensagem de erro
        db.rollback()
        return f"Erro ao apagar dados: {str(e)}"

# Função para inserir uma pessoa no banco de dados
def inserir_pessoa_bd(nome, valor_passagem, passagens):
    cursor = db.cursor()
    cursor.execute("INSERT INTO pessoas (nome, valor_passagem, passagens) VALUES (%s, %s, %s)", (nome, valor_passagem, passagens))
    db.commit()
    cursor.close()

# Rota para buscar informações do banco de dados
@app.route('/buscar_pessoas', methods=['GET'])
def buscar_pessoas():
    cursor = db.cursor()
    cursor.execute("SELECT nome, valor_passagem FROM pessoas")
    pessoas = cursor.fetchall()
    cursor.close()
    return render_template('index.html', pessoas=pessoas)


# Função para obter todas as pessoas do banco de dados
def obter_pessoas_bd():
    cursor = db.cursor()
    cursor.execute("SELECT * FROM pessoas")
    pessoas_bd = cursor.fetchall()
    cursor.close()
    return pessoas_bd

# Rota para listar todas as pessoas do banco de dados
@app.route('/listar_pessoas_bd')
def listar_pessoas_bd():
    pessoas_bd = obter_pessoas_bd()
    return render_template('pessoas_bd.html', pessoas_bd=pessoas_bd)

# Rota para inserir uma nova pessoa no banco de dados
@app.route('/inserir_pessoa_bd', methods=['POST'])
def inserir_pessoa_bd_route():
    try:
        # Obtém dados do formulário
        nome = request.form['nome']
        valor_passagem = float(request.form['valor_passagem'])
        passagens = int(request.form['passagens'])

        if valor_passagem <= 0 or passagens <= 0:
            raise ValueError("Valores inválidos. Certifique-se de inserir valores positivos.")

        # Insere a pessoa no banco de dados
        inserir_pessoa_bd(nome, valor_passagem, passagens)

        # Redireciona para a página de listagem de pessoas do banco de dados
        return redirect(url_for('listar_pessoas_bd'))

    except ValueError as e:
        # Em caso de erro de validação, exibe uma mensagem de erro na página
        return render_template('inserir.html', mensagem_erro=str(e))

@app.route('/inserir_pessoa_db', methods=['POST'])
def inserir_pessoa_db():
    cursor = db.cursor()
    nome = request.form['nome']
    valor_passagem = float(request.form['valor_passagem'])
    passagens = int(request.form['passagens'])
    cursor.execute("INSERT INTO pessoas (nome, valor_passagem, passagens) VALUES (%s, %s, %s)", (nome, valor_passagem, passagens))
    db.commit()
    cursor.close()
    return redirect(url_for('index'))

# Rota para excluir uma pessoa do banco de dados
@app.route('/excluir_pessoa_bd/<int:pessoa_id>')
def excluir_pessoa_bd(pessoa_id):
    cursor = db.cursor()
    cursor.execute("DELETE FROM pessoas WHERE id = %s", (pessoa_id,))
    db.commit()
    cursor.close()
    return redirect(url_for('listar_pessoas_bd'))

# Rota para editar uma pessoa do banco de dados
@app.route('/editar_pessoa_bd/<int:pessoa_id>', methods=['GET', 'POST'])
def editar_pessoa_bd(pessoa_id):
    if request.method == 'GET':
        cursor = db.cursor()
        cursor.execute("SELECT * FROM pessoas WHERE id = %s", (pessoa_id,))
        pessoa = cursor.fetchone()
        cursor.close()
        return render_template('editar_pessoa.html', pessoa=pessoa)
    elif request.method == 'POST':
        nome = request.form['nome']
        valor_passagem = float(request.form['valor_passagem'])
        passagens = int(request.form['passagens'])

        cursor = db.cursor()
        cursor.execute("UPDATE pessoas SET nome = %s, valor_passagem = %s, passagens = %s WHERE id = %s",
                       (nome, valor_passagem, passagens, pessoa_id))
        db.commit()
        cursor.close()
        return redirect(url_for('listar_pessoas_bd'))


# Lista para armazenar informações das pessoas
pessoas = []

@app.route('/atualizar')
def atualizar():
    # Aqui você pode adicionar a lógica necessária para renderizar a página de atualização de pessoa
    return render_template('atualizar.html')

# Função para atualizar as informações no banco de dados
def atualizar_pessoas():
    cursor = db.cursor()
    cursor.execute("SELECT * FROM pessoas")
    pessoas = cursor.fetchall()
    return pessoas

# Rota principal que renderiza a página inicial
@app.route('/')
def index():
    cursor = db.cursor()
    cursor.execute("SELECT * FROM pessoas")
    pessoas = cursor.fetchall()
    cursor.close()
    return render_template('index.html', pessoas=pessoas)

@app.route('/atualizar_pessoas', methods=['POST'])
def atualizar_pessoas_route():
    try:
        cursor = db.cursor()
        cursor.execute("UPDATE pessoas SET valor_passagem = %s, passagens = %s WHERE id = %s",
                       (request.form['valor_passagem'], request.form['passagens'], request.form['id']))
        db.commit()
        flash('Pessoa atualizada com sucesso!', 'success')
    except Exception as e:
        db.rollback()
        flash(f'Erro ao atualizar pessoa: {str(e)}', 'error')
    finally:
        cursor.close()

    return redirect(url_for('index'))


# Rota para inserir informações de uma pessoa
@app.route('/inserir_pessoa', methods=['POST'])
def inserir_pessoa():
    try:
        # Obtém dados do formulário
        nome = request.form['nome']
        valor_passagem = float(request.form['valor_passagem'])
        passagens = int(request.form['passagens'])

        if valor_passagem <= 0 or passagens <= 0:
            raise ValueError("Valores inválidos. Certifique-se de inserir valores positivos.")

        # Calcula o total do valor e a quantidade de notas e moedas necessárias
        total_valor = passagens * valor_passagem
        qtd_notas_moedas = calcular_notas_moedas(total_valor)

        # Adiciona informações à lista de pessoas
        pessoas.append((nome, valor_passagem, passagens, qtd_notas_moedas))

        # Redireciona para a página de pessoas
        return redirect(url_for('listar_pessoas'))

    except ValueError as e:
        # Em caso de erro de validação, exibe uma mensagem de erro na página
        return render_template('index.html', pessoas=pessoas, mensagem_erro=str(e))

# Rota para a página de pessoas inseridas
@app.route('/pessoas')
def listar_pessoas():
    nomes_pessoas = [pessoa[0] for pessoa in pessoas]
    return render_template('pessoas.html', nomes_pessoas=nomes_pessoas)

# Rota para a página com o total de moedas
@app.route('/total_moedas/<int:pessoa_index>')
def detalhes_moedas(pessoa_index):
    pessoa = pessoas[pessoa_index]
    total_notas_moedas = pessoa[3]
    return render_template('total_moedas.html', pessoa=pessoa, total_notas_moedas=total_notas_moedas, exibir_notas_moedas=exibir_notas_moedas)

# Rota para a página com o total de moedas para todas as pessoas
@app.route('/total_moedas_global')
def total_moedas_global():
    total_notas_moedas = calcular_total_notas_moedas(pessoas)
    detalhes_total_moedas = detalhes_notas_moedas(total_notas_moedas)
    detalhes_por_pessoa = detalhes_notas_moedas_por_pessoa([pessoa[3] for pessoa in pessoas])
    return render_template('total_moedas.html', pessoas=pessoas, total_notas_moedas=calcular_total_notas_moedas(pessoas), exibir_notas_moedas=exibir_notas_moedas,detalhes_total_moedas=detalhes_total_moedas,detalhes_por_pessoa=detalhes_por_pessoa)

# Função para calcular os detalhes das notas e moedas utilizadas por pessoa
def detalhes_notas_moedas_por_pessoa(qtd_notas_moedas):
    notas_moedas = [50, 10, 5, 2, 1, 0.5, 0.25, 0.1, 0.05, 0.01]
    detalhes_por_pessoa = []

    for i in range(len(qtd_notas_moedas)):
        detalhes = []
        for j in range(len(notas_moedas)):
            if qtd_notas_moedas[i][j] > 0:
                if notas_moedas[j] >= 1:
                    detalhes.append(f"{qtd_notas_moedas[i][j]} nota(s) de R$ {notas_moedas[j]}")
                else:
                    detalhes.append(f"{qtd_notas_moedas[i][j]} moeda(s) de R$ {notas_moedas[j]:.2f}")
        detalhes_por_pessoa.append(detalhes)

    return detalhes_por_pessoa

#
def detalhes_notas_moedas(qtd_notas_moedas):
    notas_moedas = [50, 10, 5, 2, 1, 0.5, 0.25, 0.1, 0.05, 0.01]
    detalhes = []

    for i in range(len(notas_moedas)):
        if qtd_notas_moedas[i] > 0:
            if notas_moedas[i] >= 1:
                detalhes.append(f"{qtd_notas_moedas[i]} nota(s) de R$ {notas_moedas[i]}")
            else:
                detalhes.append(f"{qtd_notas_moedas[i]} moeda(s) de R$ {notas_moedas[i]:.2f}")
    return detalhes

# Função para calcular a quantidade de notas e moedas necessárias
def calcular_notas_moedas(valor):
    notas_moedas = [50, 10, 5, 2, 1, 0.5, 0.25, 0.1, 0.05, 0.01]
    qtd_notas_moedas = [0] * len(notas_moedas)

    for i in range(len(notas_moedas)):
        qtd_notas_moedas[i] = int(valor // notas_moedas[i])
        valor %= notas_moedas[i]

    return qtd_notas_moedas

# Função para exibir as notas e moedas
def exibir_notas_moedas(qtd_notas_moedas):
    notas_moedas = [50, 10, 5, 2, 1, 0.5, 0.25, 0.1, 0.05, 0.01]
    mensagem = ""

    for i in range(len(notas_moedas)):
        if qtd_notas_moedas[i] > 0:
            if notas_moedas[i] >= 1:
                mensagem += f"{qtd_notas_moedas[i]} nota(s) de {notas_moedas[i]}\n"
            else:
                mensagem += f"{qtd_notas_moedas[i]} moeda(s) de {notas_moedas[i]:.2f}\n"

    return mensagem

# Função para calcular o total de notas e moedas utilizadas por todas as pessoas
def calcular_total_notas_moedas(pessoas):
    total_notas_moedas = [0] * 10

    for pessoa in pessoas:
        for i in range(10):
            total_notas_moedas[i] += pessoa[3][i]

    return total_notas_moedas

#
@app.route('/inserir')
def inserir():
    return render_template('inserir.html')

#
@app.route('/limpar')
def limpar():
    pessoas.clear()
    return render_template('index.html', pessoas=pessoas, total_notas_moedas=calcular_total_notas_moedas(pessoas), exibir_notas_moedas=exibir_notas_moedas)

#
@app.route('/exportar_excel')
def exportar_excel():
    if not pessoas:
        return "Nenhuma pessoa para exportar."

    # Criando um DataFrame do pandas com os dados
    df = pd.DataFrame(pessoas, columns=['Nome', 'Valor', 'Quantidade de Passagens', 'Quantidade de Moedas'])

    # Exportando para um arquivo Excel (xlsx)
    excel_filename = 'pessoas_data.xlsx'
    df.to_excel(excel_filename, index=False)

    return f'Dados exportados para {excel_filename}.'

#
@app.route('/gerar_pdf')
def gerar_pdf():
    # Crie um objeto PDF
    response = canvas.Canvas("dados_pessoas.pdf")

    # Adicione os cabeçalhos do PDF
    response.drawString(100, 800, "Nome")
    response.drawString(200, 800, "Valor")
    response.drawString(300, 800, "Qtd Passagens")
    response.drawString(400, 800, "Qtd de Moedas")

    # Adicione os dados da lista de pessoas ao PDF
    y_position = 780
    for pessoa in pessoas:
        response.drawString(100, y_position, str(pessoa[0]))
        response.drawString(200, y_position, str(pessoa[1]))
        response.drawString(300, y_position, str(pessoa[2]))
        response.drawString(400, y_position, str(pessoa[3]))

        y_position -= 20

    # Salve o PDF e finalize o objeto
    response.save()

    return "PDF gerado com sucesso!"

#
@app.route('/exportar_word')
def exportar_word():
    # Criar um novo documento Word
    doc = Document()

    # Adicionar um título ao documento
    doc.add_heading('Relatório de Pessoas', 0)

    # Adicionar uma tabela com cabeçalhos
    table = doc.add_table(rows=1, cols=4)
    table.autofit = True

    header_cells = table.rows[0].cells
    header_cells[0].text = 'Nome'
    header_cells[1].text = 'Valor'
    header_cells[2].text = 'Qtd de Passagens'
    header_cells[3].text = 'Qtd de Moedas'

    # Adicionar os dados da lista de pessoas à tabela
    for pessoa in pessoas:
        row_cells = table.add_row().cells
        row_cells[0].text = pessoa[0]  # Nome
        row_cells[1].text = str(pessoa[1])  # Valor
        row_cells[2].text = str(pessoa[2])  # Quantidade de Passagens
        row_cells[3].text = '\n'.join(exibir_notas_moedas(pessoa[3]))  # Quantidade de Moedas

    # Salvar o documento em um arquivo temporário
    temp_filename = 'relatorio_pessoas.docx'
    doc.save(temp_filename)

    # Enviar o arquivo para o cliente
    return send_file(temp_filename, as_attachment=True, download_name='relatorio_pessoas.docx')

# Executa o aplicativo se este script for executado diretamente
if __name__ == '__main__':
    app.run(debug=True)
